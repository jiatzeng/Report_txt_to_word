# -*- coding: utf8 -*-

'''
Created on 2010/1/11

@author: Jocabion
'''


import time
import re
import subprocess
import getpass
import sys
import telnetlib
import SendKeys
import socket
import random
import string   
import logging
import os,sys
import xlrd 
import fileinput
import datetime
import shutil, errno
import subprocess
import distutils.dir_util
import wmi 
import math
import _winreg
import ctypes



from threading import Thread


from datetime import datetime
from ConfigParser import SafeConfigParser
from ConfigParser import RawConfigParser
from docx import Document
from docx.shared import Inches


from subprocess import Popen,PIPE
from decimal import Decimal


if __name__ == '__main__':
    
    
    save_profile = os.getcwd() + '//Profile'
    
    save_result = os.getcwd() +'//Result//Throughput2g.txt'
    
    chrome_path = os.getcwd() + '//chromedriver_win32//chromedriver.exe'
    
    chariot_path = os.getcwd() + '/Chariot_tst'       
              
    document = Document(os.getcwd() +'\jia.docx')
    
    list =[]
    
    for table in document.tables:
     for row in table.rows:
        for cell in row.cells:
#             match_text = re.split('\n',cell.text)
            
#             print cell.text
            
            list.append(cell.text)
            
#     print list
#     sum = 0
#     for item in list :
#          
#         print sum
#         print item
#         sum = sum + 1
    
    for i in range(0,len(list),1):
      if i >20 and i < 30 :
        
#         print list[i]
        list[i] = '666'
        
#       elif i > 30 and i < 40 :
#           
#         print list[i]   
#       
#       elif i > 40 and i < 50 :
#           
#         print list[i]     
#         
#       elif i > 50 and i < 60 :
#           
#         print list[i]     
#         
#       elif i > 60 and i < 70 :
#           
#         print list[i]    
#         
      elif i > 90 and i < 100 :
           
            list[i] = '777'    
#         
#       elif i > 100 and i < 110 :
#           
#         print list[i]     
#         
#       elif i > 110 and i < 120 :
#           
#         print list[i]   
#         
#       elif i > 120 and i < 130 :
#           
#         print list[i]   
#       
#       elif i > 130 and i < 140 :
#           
#         print list[i]   
#         
#       elif i > 176 and i < 180 :
#           
#         print list[i]    
#       
#       elif i > 180 and i < 184 :
#           
#         print list[i]  
#       
#       elif i > 184 and i < 188 :
#           
#         print list[i]    
#             
#       elif i > 188 and i < 192 :
#           
#         print list[i]    
#       
#       elif i > 192 and i < 196 :
#           
#         print list[i]  
#       
      elif i > 196 and i < 200 :
           
          list[i] = '888'      
        
#     print list
    
    
    
#     for table in document.tables:
#      for row in table.rows:
#        for cell in row.cells:   
    print len(document.tables)
    print len(document.tables[0].rows)
    print len(document.tables[0].rows[0].cells)
    
    sum = 0
    l = 0

    for k in range(0,len(document.tables),1):
     for i in range(0,len(document.tables[l].rows),1):    
      for j in range(0,len(document.tables[l].rows[l].cells),1): 
#            for i in range(0,len(list),1): 
               
               print document.tables[k].rows[i].cells[j].text 
               document.tables[k].rows[i].cells[j].text = list[sum]
               sum = sum + 1
               print sum
  
     l = l+1
  
  
  
  
                 
#     for j in range(0,10,1): 
# 
#                print document.tables[0].rows[1].cells[j].text 
               
               
               
               

#     print row.cells[j].text
    
    
    
    
    
    
    document.save('jia.docx')
        
#         print list.index(str(item))
#         elif item == 'WEP-64' :  
#             print list.index(str(item))  
#         elif item == 'WEP-128' :
#             print list.index(str(item)) 
#         elif item == 'WPA-TKIP' :    
#             print list.index(str(item)) 
#         elif item == 'WPA2-AES' :   
#             print list.index(str(item)) 
 
                
            
            
        
    
            
#             match_text = re.findall('[^a-zA-z]',cell.text)
#             sum =''
#             for i in range(0,len(match_text),1) :
#  
#                sum = sum + str(match_text[i])
#                print sum
    


